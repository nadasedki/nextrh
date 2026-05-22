import sys
import json
import os
from docx import Document

def smart_replace(docx_path, mapping_file):
    with open(mapping_file, 'r', encoding='utf-8') as f:
        mapping = json.load(f)

    doc = Document(docx_path)
    
    # On trie du plus long au plus court
    sorted_mapping = dict(sorted(mapping.items(), key=lambda item: len(str(item[0])), reverse=True))

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for old_text, new_text in sorted_mapping.items():
                if str(old_text).strip() and str(old_text).strip() in p.text:
                    # LOGIQUE : Si on trouve la clé, on VIDE le paragraphe et on met la valeur
                    # Cela évite de laisser des morceaux du template original
                    p.text = str(new_text) 
                    break # On passe au paragraphe suivant dès qu'un match est fait

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    doc.save(docx_path)

if __name__ == "__main__":
    try:
        smart_replace(sys.argv[1], sys.argv[2])
        print("PYTHON_SUCCESS")
    except Exception as e:
        print(f"PYTHON_ERROR: {str(e)}")
        sys.exit(1)